from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build_architecture_diagram(output_path: Path):
    width, height = 1400, 820
    image = Image.new("RGB", (width, height), color=(248, 250, 255))
    draw = ImageDraw.Draw(image)

    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/segoeuib.ttf", 36)
        font_box_title = ImageFont.truetype("C:/Windows/Fonts/segoeuib.ttf", 24)
        font_box = ImageFont.truetype("C:/Windows/Fonts/segoeui.ttf", 18)
    except Exception:
        font_title = ImageFont.load_default()
        font_box_title = ImageFont.load_default()
        font_box = ImageFont.load_default()

    draw.text((40, 20), "Speech Disability Interpreter - Business Architecture", fill=(22, 35, 68), font=font_title)

    boxes = [
        {
            "title": "Channel Layer",
            "body": "Web UI\nAPI Clients\nBatch Jobs",
            "xy": (70, 120, 360, 320),
            "fill": (231, 244, 255),
        },
        {
            "title": "API and Orchestration",
            "body": "FastAPI Service\nRequest Routing\nProcessing Control",
            "xy": (430, 120, 770, 320),
            "fill": (238, 251, 241),
        },
        {
            "title": "Interpretation Pipeline",
            "body": "Audio Preprocessing\nSpeech Recognition\nAdaptive Correction\nText-to-Speech",
            "xy": (840, 120, 1270, 320),
            "fill": (255, 245, 233),
        },
        {
            "title": "Personalization and Learning",
            "body": "User Profiles\nFeedback Capture\nPattern Memory",
            "xy": (250, 420, 620, 670),
            "fill": (243, 241, 255),
        },
        {
            "title": "Data and Outputs",
            "body": "Audio Storage\nCorrected Text\nSynthesized Audio\nOperational Logs",
            "xy": (690, 420, 1090, 670),
            "fill": (255, 239, 242),
        },
    ]

    def draw_box(box):
        x1, y1, x2, y2 = box["xy"]
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill=box["fill"], outline=(80, 98, 132), width=3)
        draw.text((x1 + 20, y1 + 18), box["title"], fill=(25, 40, 78), font=font_box_title)

        text_y = y1 + 64
        for line in box["body"].split("\n"):
            draw.text((x1 + 24, text_y), f"- {line}", fill=(48, 63, 96), font=font_box)
            text_y += 34

    for box in boxes:
        draw_box(box)

    def arrow(start, end):
        draw.line([start, end], fill=(74, 98, 142), width=5)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 16, ey - 10), (ex - 16, ey + 10)], fill=(74, 98, 142))

    arrow((360, 220), (430, 220))
    arrow((770, 220), (840, 220))
    arrow((550, 320), (550, 420))
    arrow((960, 320), (960, 420))
    arrow((620, 545), (690, 545))

    draw.text((70, 740), "Flow: Channel -> API -> Interpretation -> Personalization/Data -> User-facing outputs", fill=(64, 78, 108), font=font_box)
    image.save(output_path)


def main():
    base_dir = Path(__file__).resolve().parents[1]
    output_dir = base_dir / "demo"
    output_dir.mkdir(parents=True, exist_ok=True)
    architecture_diagram = output_dir / "architecture_diagram.png"
    build_architecture_diagram(architecture_diagram)

    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run("AI Speech Disability Interpreter\nBusiness Overview")
    run.bold = True
    run.font.size = Pt(24)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph("Prepared for stakeholder and competition presentation | June 2026")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    add_heading(doc, "1. Executive Summary", level=1)
    doc.add_paragraph(
        "The AI Speech Disability Interpreter is an assistive communication platform designed to improve understanding "
        "of speech from individuals with dysarthria and related speech impairments. The solution converts impaired speech "
        "into clear text and intelligible synthesized audio, enabling smoother communication in healthcare, daily life, "
        "education, and digital services."
    )
    add_bullets(doc, [
        "Mission: reduce communication barriers for people with speech disabilities.",
        "Outcome: faster and clearer communication between speakers and listeners.",
        "Positioning: practical assistive AI solution with personalization and measurable performance."
    ])

    add_heading(doc, "2. Business Problem and Opportunity", level=1)
    doc.add_paragraph(
        "People with motor speech disorders are frequently misunderstood by conventional voice interfaces and by "
        "unfamiliar listeners. This creates friction in critical interactions such as clinical appointments, customer service, "
        "and public-facing communication."
    )
    add_bullets(doc, [
        "Current alternatives are either generic ASR tools or high-friction manual communication methods.",
        "There is a significant unmet accessibility need for reliable, real-time interpretation support.",
        "Organizations increasingly prioritize inclusive technology and measurable accessibility outcomes."
    ])

    add_heading(doc, "3. Solution Overview", level=1)
    doc.add_paragraph(
        "The platform delivers a full interpretation pipeline: audio intake, speech transcription, correction of impaired "
        "patterns, and speech regeneration. It also includes user-level personalization to improve outcomes over time."
    )
    add_numbered(doc, [
        "Audio preprocessing to normalize and improve input quality.",
        "Automatic speech recognition to generate raw text.",
        "Adaptive correction pipeline to improve readability and clarity.",
        "Personalization layer to learn user-specific patterns.",
        "Text-to-speech generation of corrected output."
    ])

    add_heading(doc, "4. Architecture (Business View)", level=1)
    doc.add_paragraph(
        "The system is designed as a modular service architecture to support fast iteration, controlled scaling, and "
        "future enterprise integration."
    )
    add_bullets(doc, [
        "API Layer: FastAPI endpoints for transcription, feedback, batch processing, and output retrieval.",
        "Inference Layer: speech recognition and correction engines for interpretation quality.",
        "Personalization Layer: user profile management and correction learning.",
        "Output Layer: text response plus synthesized speech output.",
        "Observability Layer: stage-level timing and runtime telemetry for performance reporting."
    ])

    diagram_title = doc.add_paragraph("Business Architecture Diagram")
    diagram_title.runs[0].bold = True
    doc.add_picture(str(architecture_diagram), width=Pt(470))
    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_heading(doc, "5. Key Features", level=1)
    add_bullets(doc, [
        "Real-time audio transcription with confidence scoring.",
        "Intelligent correction combining model-based and rule-based processing.",
        "Adaptive correction strategy based on confidence and speech characteristics.",
        "Personalized correction memory per user.",
        "Corrected speech audio generation.",
        "Batch processing for operational workflows.",
        "Demo dashboard with timing visualization and explainable outputs."
    ])

    add_heading(doc, "6. Value Proposition", level=1)
    add_bullets(doc, [
        "Accessibility impact: improves inclusivity and communication equity.",
        "Operational impact: reduces repeated clarification in high-touch environments.",
        "User impact: enables higher confidence and lower social friction for speakers.",
        "Business impact: creates a differentiated assistive AI product with clear measurable outcomes."
    ])

    add_heading(doc, "7. Target Use Cases", level=1)
    add_bullets(doc, [
        "Healthcare interactions between patients and providers.",
        "Customer support and front-desk communication.",
        "Education and student support services.",
        "Public service counters and accessibility programs.",
        "Personal daily communication support through companion apps."
    ])

    add_heading(doc, "8. Product Demonstration Capability", level=1)
    doc.add_paragraph("The current implementation supports a complete demonstration workflow:")
    add_numbered(doc, [
        "Upload impaired speech sample.",
        "Display original and corrected text output.",
        "Show correction strategy and confidence.",
        "Play corrected synthesized speech.",
        "Review timing and performance metrics.",
        "Trigger batch mode and user statistics for extended scenario coverage."
    ])

    add_heading(doc, "9. Governance and Security Posture", level=1)
    doc.add_paragraph(
        "The current version is optimized for demo and pilot environments. For production deployment, the following "
        "controls are recommended as standard enterprise hardening steps."
    )
    add_bullets(doc, [
        "API authentication and role-based authorization.",
        "File upload validation and request size controls.",
        "Audit logging, retention policy, and monitoring.",
        "Data privacy controls for stored audio and profile records.",
        "Environment-based secret management and secure configuration."
    ])

    add_heading(doc, "10. Implementation Status", level=1)
    add_bullets(doc, [
        "Core backend pipeline implemented and operational.",
        "Frontend demo console implemented for live walkthrough.",
        "Demo media package available (video + presentation deck).",
        "Launch scripts available for local startup and demonstration."
    ])

    add_heading(doc, "11. Recommended Next Steps", level=1)
    add_numbered(doc, [
        "Package pilot-ready release with configuration profiles.",
        "Introduce authentication and policy-based access controls.",
        "Create KPI dashboard for business reporting (accuracy, latency, user impact).",
        "Run controlled pilot with representative user group and collect outcomes.",
        "Prepare commercialization plan: deployment model, support model, and pricing structure."
    ])

    add_heading(doc, "12. Conclusion", level=1)
    doc.add_paragraph(
        "This solution demonstrates a strong foundation for an assistive communication product with clear social value "
        "and practical business applicability. It is well-suited for stakeholder demos, pilot engagements, and further "
        "productization planning."
    )

    output_path = output_dir / "Business_Overview_Speech_Disability_Interpreter.docx"
    doc.save(output_path)
    print(f"Document created: {output_path}")


if __name__ == "__main__":
    main()
